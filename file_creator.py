from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from io import BytesIO

class Create_Doc:
    def __init__(self) -> None:
        self.doc = Document()
    
    def markdown_to_word(self,markdown_text):
        # Convert Markdown to HTML
        # html = markdown.markdown(markdown_text)

        # Parse the Markdown text and add formatted content to the document
        for line in markdown_text.split('\n'):
            if line.startswith('# '):
                heading = line[2:]
                p = self.doc.add_heading(heading, level=1)
            elif line.startswith('## '):
                heading = line[3:]
                p = self.doc.add_heading(heading, level=2)
            elif line.startswith('### '):
                heading = line[4:]
                p = self.doc.add_heading(heading, level=3)
            elif line.startswith('- '):
                item = line[2:]
                p = self.doc.add_paragraph(item, style='ListBullet')
            else:
                p = self.doc.add_paragraph(line)
            
            # Adjust paragraph formatting (optional)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.style.font.size = Pt(12)

        # Save the document to a BytesIO object
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer